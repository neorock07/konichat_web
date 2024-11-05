import streamlit as st
from controller.document_controller import (
    delete_docs, upload_document, update_document,
    get_all_docs, get_feedback_doc_by_id, 
    upload_document_feedback, delete_feedback_docs
)
from controller.auth_controller import get_all_role
from controller.feedback_method import (
    get_count_feedback, get_bad_feedback
)
from datetime import datetime
import docx
from io import BytesIO

st.set_page_config("KoniChat | Admin", page_icon="assets/fav.png")

# Authentication check
if 'user_data' not in st.session_state:
    st.switch_page("login.py")

st.header("Chatbot Performance")

# Load role and document data
role = get_all_role()
data_count_feedback = get_count_feedback()
data_docs = get_all_docs()
list_role = [i['name'] for i in role]
roles_feed = [{i['name']: i['id']} for i in role]
id_role_selected = None

# Sidebar - Upload new document
if role != "error" and role is not None:
    with st.sidebar:
        st.header("Upload New Doc.")
        st.warning("Upload dokumen sesuai role!")
        role_select = st.selectbox("Role", options=list_role)
        file = st.file_uploader("Dokumen", type="pdf", accept_multiple_files=True)
        
        for i in role:
            if i['name'] == role_select:
                id_role_selected = i['id']

        if st.button(f"Upload {role_select}"):
            if file:
                for i in file:
                    upload_document(i, id_role_selected)
            else:
                st.toast("Please provide file!")
            st.rerun()

def update_doc(file, id_role):
    if file:
        formatted_time = datetime.today().strftime("%Y-%m-%d %H:%M:%S")
        update_document(file, id_role, formatted_time)
    else:
        st.toast("Please provide file!")

# Chart - Feedback count
if data_count_feedback != "error":
    st.bar_chart(data_count_feedback, x="tipe", y="count")

########################################################################
##                     Dokumen Sumber                                 ##
########################################################################
@st.fragment
@st.dialog(title="Yakin ingin menghapus?")
def dialog_konf_doc(id_doc):
    if st.button("Yakin"):
        delete_docs(id_doc)
        st.rerun()

@st.fragment
def widget_doc_sumber():
    st.divider()
    st.header("Update Document Sumber")
    
    group_doc = {}
    for i in data_docs:
        name = i['name']
        if name not in group_doc:
            group_doc[name] = []
        group_doc[name].append(i)

    roles = ['--pilih--'] + list(group_doc.keys())
    selected_role = st.selectbox("Role", options=roles)
    selected_doc, selected_index = None, None
    
    if selected_role != '--pilih--':
        docs_title = [i['title'] for i in group_doc[selected_role]]
        docs_id = group_doc[selected_role]
        selected_doc = st.selectbox("Dokumen", options=docs_title, placeholder="search doc...")
        selected_index = docs_title.index(selected_doc) if selected_doc else None
    else:
        st.write("Dokumen")
        st.warning("Pilih role yang valid!")

    if selected_index is not None:
        id_doc_update = docs_id[selected_index]['id']
        file_update = st.file_uploader("Upload updated document", type=["pdf"])
        
        col1, col2 = st.columns([0.2, 0.8])
        with col1:
            if st.button(f"✏ Update {selected_role}"):
                update_doc(file_update, id_doc_update)
                st.rerun()
        with col2:
            if st.button(f"🗑 Delete {selected_role}", type="primary"):
                dialog_konf_doc(id_doc_update)

widget_doc_sumber()

########################################################################
##                     Dokumen Feedback                               ##
########################################################################

@st.fragment
@st.dialog(title="Yakin ingin menghapus?")
def dialog_konf_feedback(id_role):
    if st.button("Yakin"):
        delete_feedback_docs(id_role)
        st.rerun()

@st.fragment
def widget_doc_feedback(roles):
    st.divider()
    st.title("Document Feedback")

    opt = [next(iter(val)) for val in roles]
    selected_role_feed = st.selectbox("Role", options=opt, key="feedback_select")

    if selected_role_feed != '--pilih--':
        id_role_to_feed = next(i[selected_role_feed] for i in roles if next(iter(i)) == selected_role_feed)
        doc_feeds = get_feedback_doc_by_id(id_role_to_feed)
        
        if doc_feeds:
            st.caption(f"Nama file: {doc_feeds['title']}")
        else:
            st.warning("⚠ Data Masih Kosong! Silahkan Tambahkan")
        # """
        #     membuat file words dari kumpulan data feedback bertipe bad;
        #     sebagai template untuk dokumen revisi feedback
        # """
        docs = docx.Document()
        bad_data = get_bad_feedback(id_role=id_role_to_feed)

        if not isinstance(bad_data, str):
            for i in range(len(bad_data)):
                docs.add_paragraph("").add_run(f"{i+1}. Query : ").bold = True
                docs.add_paragraph(bad_data[i]['human_query'])
                docs.add_paragraph("").add_run(f"Your Rejected Response : ").bold = True
                docs.add_paragraph(bad_data[i]['ai_respon'])
                docs.add_paragraph("").add_run(f"Your Chosen Response : ").bold = True
                docs.add_paragraph("<EOS>")

            buffer = BytesIO()
            docs.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="Unduh feedback user",
                data=buffer,
                file_name=f"dokumen_revisi_{selected_role_feed}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        feedback_doc = st.file_uploader("Upload feedback document", type="pdf")
        
        col1, col2 = st.columns([0.2, 0.8])
        # """
        #     tombol untuk upload dokumen revisi
        # """
        with col1:
            if st.button(f"✏ Upload {selected_role_feed}"):
                if feedback_doc:
                    upload_document_feedback(feedback_doc, id_role_to_feed)
                    st.rerun(scope=["fragment"])
                else:
                    st.toast("❗ Please select a document")
        # """
        #     tombol untuk delete dokumen revisi
        # """
        with col2:
            if st.button(f"🗑 Delete {selected_role_feed}", type="primary", key="awasadasule"):
                dialog_konf_feedback(id_role_to_feed)
    else:
        st.warning("Pilih role yang valid!")

widget_doc_feedback(roles_feed)
